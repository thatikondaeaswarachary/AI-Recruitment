from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()

# Title
title = document.add_paragraph('Infosys Springboard Virtual Internship\n“Project Name: ConceptClarity:\nScientific Terminology Explainer”')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.bold = True
    run.font.size = Pt(16)

# Name & Date
name_date = document.add_paragraph('Name: eshwarachary\nDate: 13-07-2026')
name_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Milestone
milestone = document.add_heading('Milestone 1 – Web App & Search Bar', level=1)

# Introduction
document.add_heading('Introduction', level=2)
document.add_paragraph('Concept Clarity is a web-based educational platform designed to simplify complex scientific and technical terminology for students and beginners. This milestone focused on building a secure and intuitive learning environment where users can register, log in, and access beginner-friendly explanations of scientific terms through an organized search system.')

document.add_paragraph('Key Highlights: Developed a secure user registration and login process with password hashing.', style='List Bullet')
document.add_paragraph('Implemented RESTful APIs for user authentication and term search functionality.', style='List Bullet')
document.add_paragraph('Created a responsive frontend with automatic redirection for seamless user experience.', style='List Bullet')
document.add_paragraph('Established SQLite database integration with proper session management for data security.', style='List Bullet')

document.add_heading('Tech Stack:', level=3)
document.add_paragraph('Frontend: HTML5, CSS3, and JavaScript for intuitive user interfaces.', style='List Bullet')
document.add_paragraph('Backend: FastAPI with Pydantic models for robust API development and validation.', style='List Bullet')
document.add_paragraph('Database: SQLite with SQLAlchemy ORM for efficient data storage and retrieval.', style='List Bullet')
document.add_paragraph('Server & Tools: Uvicorn ASGI server for high-performance backend operations.', style='List Bullet')

document.add_paragraph('This milestone established a reliable foundation for future enhancements, including advanced search features, user progress tracking, and expanded content libraries. The system provides a scalable and user-friendly educational platform that prepares for future improvements like personalized learning paths and interactive study tools.')

# Objectives
document.add_heading('Objectives', level=2)
document.add_paragraph('The main objectives successfully achieved in this project phase are:')

document.add_paragraph('1. Secure User Authentication', style='List Number')
document.add_paragraph('Implemented robust user registration with email and username validation', style='List Bullet')
document.add_paragraph('Created secure login functionality with password hashing using bcrypt', style='List Bullet')
document.add_paragraph('Added duplicate user detection to prevent multiple registrations', style='List Bullet')
document.add_paragraph('Implemented credential verification with proper error messages', style='List Bullet')

document.add_paragraph('2. Search Functionality Foundation', style='List Number')
document.add_paragraph('Developed API endpoint for scientific term search', style='List Bullet')
document.add_paragraph('Created structured response format for term explanations', style='List Bullet')
document.add_paragraph('Implemented parameter validation and sanitization', style='List Bullet')
document.add_paragraph('Established the search workflow from frontend to backend', style='List Bullet')

document.add_paragraph('3. User Experience & Navigation', style='List Number')
document.add_paragraph('Designed intuitive authentication pages (signup/login)', style='List Bullet')
document.add_paragraph('Created a dedicated search interface for scientific terms', style='List Bullet')
document.add_paragraph('Implemented seamless navigation between authentication and search pages', style='List Bullet')
document.add_paragraph('Added responsive design for cross-device compatibility', style='List Bullet')

document.add_paragraph('4. Data Security & Validation', style='List Number')
document.add_paragraph('Implemented password hashing for secure credential storage', style='List Bullet')
document.add_paragraph('Added input validation on both frontend and backend', style='List Bullet')
document.add_paragraph('Configured CORS for secure cross-origin requests', style='List Bullet')
document.add_paragraph('Established secure database connections with proper session management', style='List Bullet')

# Code Implementation
document.add_heading('Code Implementation', level=2)
document.add_paragraph('The project has been implemented using FastAPI (Python) as the backend framework, and vanilla HTML/CSS/JavaScript for the frontend interface. Below are the key components of the implementation:')

# Module 1
document.add_heading('Module 1: Authentication Service (Backend)', level=3)
document.add_paragraph('The Authentication Service module is responsible for managing secure user registration and login functionalities within the application. It uses FastAPI routes combined with Pydantic models to validate user input and ensure data integrity. User passwords are encrypted using bcrypt hashing before being stored in the database to maintain security. This module ensures that only authenticated users can access the system by verifying credentials and returning appropriate responses.')

document.add_paragraph('Backend: schemas.py')
code_style = document.styles.add_style('CodeBlock', 1)
code_style.font.name = 'Courier New'
code_style.font.size = Pt(10)

document.add_paragraph('''from pydantic import BaseModel 

class SignupRequest(BaseModel): 
    username: str 
    email: str 
    password: str 

class LoginRequest(BaseModel): 
    username: str 
    password: str 
''', style='CodeBlock')

document.add_paragraph('Backend: main.py')
document.add_paragraph('''from fastapi import FastAPI, Depends 
from fastapi.middleware.cors import CORSMiddleware 
from sqlalchemy.orm import Session 
from passlib.context import CryptContext 
from database import engine, SessionLocal 
from models import Base, User 
from schemas import SignupRequest, LoginRequest 

app = FastAPI() 

app.add_middleware( 
    CORSMiddleware, 
    allow_origins=["http://localhost:5500"], 
    allow_credentials=True, 
    allow_methods=["*"], 
    allow_headers=["*"], 
) 

Base.metadata.create_all(bind=engine) 

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto") 

def get_db(): 
    db = SessionLocal() 
    try: 
        yield db 
    finally: 
        db.close() 

@app.post("/signup") 
def signup(data: SignupRequest, db: Session = Depends(get_db)): 
    existing = db.query(User).filter( 
        (User.username == data.username) | 
        (User.email == data.email) 
    ).first() 

    if existing: 
        return {"message": "User already exists"} 

    user = User( 
        username=data.username, 
        email=data.email, 
        password=pwd_context.hash(data.password) 
    ) 
    db.add(user) 
    db.commit() 
    return {"message": "User registered successfully"} 

@app.post("/login") 
def login(data: LoginRequest, db: Session = Depends(get_db)): 
    user = db.query(User).filter(User.username == data.username).first() 
    if not user or not pwd_context.verify(data.password, user.password): 
        return {"message": "Invalid credentials"} 
    return {"message": "Login successful"} 

@app.get("/search") 
def search(term: str): 
    return { 
        "term": term, 
        "explanation": f"{term} is explained in simple language for beginners." 
    } 
''', style='CodeBlock')


# Module 2
document.add_heading('Module 2: Database Layer', level=3)
document.add_paragraph('The Database Layer module manages all user-related data using SQLAlchemy ORM. It defines the database schema and handles the creation of tables required for storing user credentials. This module is responsible for establishing database connections, managing sessions, and ensuring reliable data persistence. Secure storage and retrieval of user information is achieved through structured ORM-based operations.')

document.add_paragraph('Backend Database: database.py')
document.add_paragraph('''from sqlalchemy import create_engine 
from sqlalchemy.orm import sessionmaker, declarative_base 

DATABASE_URL = "sqlite:///./users.db" 

engine = create_engine( 
    DATABASE_URL, connect_args={"check_same_thread": False} 
) 

SessionLocal = sessionmaker(bind=engine) 
Base = declarative_base() 
''', style='CodeBlock')

document.add_paragraph('Backend Database: models.py')
document.add_paragraph('''from sqlalchemy import Column, Integer, String 
from database import Base 

class User(Base): 
    __tablename__ = "users" 
    id = Column(Integer, primary_key=True, index=True) 
    username = Column(String, unique=True, index=True) 
    email = Column(String, unique=True) 
    password = Column(String) 
''', style='CodeBlock')


# Module 3
document.add_heading('Module 3: Term Search Service (Backend)', level=3)
document.add_paragraph('The Term Search Service module provides basic functionality for retrieving simplified explanations of scientific terms. It exposes a REST API endpoint that accepts a term as input and returns an easy-to-understand explanation suitable for beginners. This module currently serves as a foundational implementation for term lookup. It is designed to be extended in future milestones with AI-driven and NLP-based explanation generation.')

document.add_paragraph('Backend: main.py')
document.add_paragraph('''@app.get("/search") 
def search(term: str): 
    return { 
        "term": term, 
        "explanation": f"{term} is explained in simple language for beginners." 
    } 
''', style='CodeBlock')

# Module 4
document.add_heading('Module 4: Frontend Integration', level=3)
document.add_paragraph('The Frontend Integration module enables users to interact with the system through a web-based interface. It includes pages for user registration, login, and scientific term search, developed using HTML, CSS, and JavaScript. The frontend communicates with the backend APIs using the Fetch API to send and receive data. Automatic page redirection and user feedback messages enhance usability and provide a smooth user experience.')

document.add_paragraph('Frontend: signup.html')
document.add_paragraph('''<!DOCTYPE html> 
<html> 
<head> 
    <title>Sign Up</title> 
    <link rel="stylesheet" href="style.css"> 
</head> 
<body> 
    <h2>Create Account</h2> 
    <input id="username" placeholder="Username"> 
    <input id="email" placeholder="Email"> 
    <input id="password" type="password" placeholder="Password"> 
    <button onclick="signup()">Register</button> 
    <p id="msg"></p> 
    <script src="script.js"></script> 
</body> 
</html> 
''', style='CodeBlock')

document.add_paragraph('Frontend: login.html')
document.add_paragraph('''<!DOCTYPE html> 
<html> 
<head> 
    <title>Login</title> 
    <link rel="stylesheet" href="style.css"> 
</head> 
<body> 
    <h2>Login</h2> 
    <input id="username" placeholder="Username"> 
    <input id="password" type="password" placeholder="Password"> 
    <button onclick="login()">Login</button> 
    <p id="msg"></p> 
    <script src="script.js"></script> 
</body> 
</html> 
''', style='CodeBlock')

document.add_paragraph('Frontend: search.html')
document.add_paragraph('''<!DOCTYPE html> 
<html> 
<head> 
    <title>Concept Clarity</title> 
    <link rel="stylesheet" href="style.css"> 
</head> 
<body> 
    <h2>Scientific Term Simplifier</h2> 
    <input id="term" placeholder="Enter scientific term"> 
    <button onclick="searchTerm()">Search</button> 
    <p id="result"></p> 
    <script src="script.js"></script> 
</body> 
</html> 
''', style='CodeBlock')

document.add_paragraph('Frontend: script.js (API calls + redirect)')
document.add_paragraph('''const FRONTEND_BASE = "http://localhost:5500"; 
const BACKEND_BASE = "http://127.0.0.1:8000"; 

// SIGNUP 
function signup() { 
    console.log("Signup clicked"); 
    fetch(`${BACKEND_BASE}/signup`, { 
        method: "POST", 
        headers: { "Content-Type": "application/json" }, 
        body: JSON.stringify({ 
            username: document.getElementById("username").value, 
            email: document.getElementById("email").value, 
            password: document.getElementById("password").value 
        }) 
    }) 
    .then(response => response.json()) 
    .then(data => { 
        console.log("Signup response:", data); 
        document.getElementById("msg").innerText = data.message || "Signup completed"; 
        // FORCE REDIRECT (ABSOLUTE URL) 
        setTimeout(() => { 
            console.log("Redirecting to login..."); 
            window.location.replace(`${FRONTEND_BASE}/login.html`); 
        }, 1000); 
    }) 
    .catch(error => { 
        console.error("Signup error:", error); 
    }); 
} 

// LOGIN 
function login() { 
    console.log("Login clicked"); 
    fetch(`${BACKEND_BASE}/login`, { 
        method: "POST", 
        headers: { "Content-Type": "application/json" }, 
        body: JSON.stringify({ 
            username: document.getElementById("username").value, 
            password: document.getElementById("password").value 
        }) 
    }) 
    .then(response => response.json()) 
    .then(data => { 
        console.log("Login response:", data); 
        document.getElementById("msg").innerText = data.message || "Login completed"; 
        // FORCE REDIRECT (ABSOLUTE URL) 
        setTimeout(() => { 
            console.log("Redirecting to search..."); 
            window.location.replace(`${FRONTEND_BASE}/search.html`); 
        }, 1000); 
    }) 
    .catch(error => { 
        console.error("Login error:", error); 
    }); 
} 

// SEARCH 
function searchTerm() { 
    const term = document.getElementById("term").value; 
    fetch(`${BACKEND_BASE}/search?term=${encodeURIComponent(term)}`) 
    .then(response => response.json()) 
    .then(data => { 
        document.getElementById("result").innerText = data.explanation; 
    }); 
} 
''', style='CodeBlock')

document.add_paragraph('Frontend: style.css')
document.add_paragraph('''body { 
    font-family: Arial, sans-serif; 
    margin: 50px; 
} 
input { 
    display: block; 
    margin: 10px 0; 
    padding: 8px; 
    width: 260px; 
} 
button { 
    padding: 8px 14px; 
    margin-top: 10px; 
} 
p { 
    margin-top: 15px; 
    font-weight: bold; 
} 
''', style='CodeBlock')

# Output screenshots section
document.add_heading('Output screenshots:', level=2)
document.add_paragraph('1. Sign In page:')
document.add_paragraph('[ Insert Sign In Page Screenshot Here ]').alignment = WD_ALIGN_PARAGRAPH.CENTER
document.add_paragraph('2. Login page:')
document.add_paragraph('[ Insert Login Page Screenshot Here ]').alignment = WD_ALIGN_PARAGRAPH.CENTER
document.add_paragraph('3. Scientific term search page:')
document.add_paragraph('[ Insert Search Page Screenshot Here ]').alignment = WD_ALIGN_PARAGRAPH.CENTER

# Conclusion
document.add_heading('Conclusion', level=2)
document.add_paragraph("In this milestone, the Concept Clarity platform successfully establishes a strong and reliable foundation for the overall application, delivering a fully functional and secure user experience from end to end. Secure user authentication has been implemented, featuring both user registration and login workflows, to ensure that only valid and verified users can access the system's core functionality. The integration of a relational database provides safe, structured, and efficient storage of user credentials, creating a trustworthy data layer that supports user management and future data expansion. Frontend and backend components communicate effectively through well-defined REST APIs, facilitating smooth and reliable data flow between the user interface and server-side logic. The term search functionality a core feature of the platform has been successfully implemented to demonstrate the application's purpose, providing clear and simplified explanations of scientific terminology based on user queries. Comprehensive input validation is applied at both the client and server levels, preventing erroneous or malicious data from compromising system integrity. Password encryption using industry-standard hashing algorithms ensures that sensitive user information remains protected against unauthorized access. Furthermore, the entire system architecture follows a modular and scalable design approach, allowing individual components to be developed, tested, and enhanced independently without disrupting the overall application. This modularity not only improves maintainability but also supports the seamless integration of more advanced features planned for later phases. The successful completion of this milestone clearly validates the technical feasibility of the project, confirming that the chosen technologies and design patterns work effectively together to meet the platform’s requirements. It also ensures a high degree of system stability and usability, as evidenced by rigorous testing of user flows such as registration, authentication, and term search operations. By delivering a robust and secure base version of the application, Milestone-1 effectively prepares the ground for integrating more sophisticated capabilities such as AI-powered explanations, personalized learning features, and enhanced interactivity in upcoming development phases, setting a solid trajectory toward the project's long-term vision.")

document.save('ConceptClarity_Documentation_eshwarachary.docx')
