from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

document = Document()

# Title
title = document.add_paragraph('Project Name: AI Recruitment\n')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(16)
    run.bold = True

# Name and Date
info = document.add_paragraph('Name: eshwarachary\nDate: 13-07-2026\n')
info.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Introduction
doc_heading = document.add_heading('Project Documentation', level=1)
doc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_heading('Introduction', level=2)
document.add_paragraph('The AI Recruitment platform is a full-stack web application designed to automate resume parsing, candidate-job matching, and recruiter assistance using AI. This platform streamlines the hiring process by leveraging AI models to extract candidate profiles and match them with job descriptions.')

# System Architecture
document.add_heading('System Architecture', level=2)
document.add_paragraph('1. Frontend: React.js (Vite) with TailwindCSS for styling and Recharts for analytics dashboard.')
document.add_paragraph('2. Backend: FastAPI providing RESTful APIs for the application logic.')
document.add_paragraph('3. Database: PostgreSQL/SQLite for data persistence using SQLAlchemy ORM.')
document.add_paragraph('4. AI Integration: NLP models for resume parsing and compatibility matching.')

def add_code_section(title, files):
    document.add_heading(title, level=2)
    for file_path in files:
        if os.path.exists(file_path):
            document.add_heading(os.path.basename(file_path), level=3)
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    code_content = f.read()
                    
                # Limit code content to avoid excessively long documents
                lines = code_content.split('\n')
                if len(lines) > 200:
                    code_content = '\n'.join(lines[:200]) + '\n... (truncated for brevity)'
                    
                p = document.add_paragraph(code_content)
                p.style = 'Normal'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
            except Exception as e:
                print(f"Failed to read {file_path}: {e}")
        else:
            print(f"File not found: {file_path}")

# Backend Code
backend_files = [
    'backend/main.py',
    'backend/database.py',
    'backend/models.py',
    'backend/schemas.py',
    'backend/routers/resumes.py',
]
add_code_section('Backend Implementation', backend_files)

# Frontend Code
frontend_files = [
    'frontend/src/App.jsx',
    'frontend/src/main.jsx',
    'frontend/src/components/Navbar.jsx',
    'frontend/src/pages/Dashboard.jsx',
    'frontend/src/pages/UploadResume.jsx',
    'frontend/src/pages/JobDescription.jsx',
    'frontend/src/index.css'
]
add_code_section('Frontend Implementation', frontend_files)

# Screenshots
document.add_heading('Screenshots', level=2)
document.add_paragraph('Dashboard View:')
document.add_paragraph('[INSERT DASHBOARD SCREENSHOT HERE]')
document.add_paragraph('Upload Resume View:')
document.add_paragraph('[INSERT UPLOAD RESUME SCREENSHOT HERE]')
document.add_paragraph('Job Matching View:')
document.add_paragraph('[INSERT MATCHING SCREENSHOT HERE]')

document.save('AI_Recruitment_Documentation_eshwarachary.docx')
